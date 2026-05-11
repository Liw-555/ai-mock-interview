"""
Resume parser: detect format, extract text, and structure via LLM.
"""

import os
from typing import Tuple, Optional, Dict, Any

from models.database import FileType


def parse_resume_file(file_path: str, file_type: FileType) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
    """
    Parse a resume file and return raw_text + structured content_json.
    Returns (raw_text, content_json) or raises Exception on failure.
    """
    # Step 1: Extract raw text based on file type
    if file_type == FileType.pdf:
        raw_text = extract_pdf_text(file_path)
    elif file_type == FileType.docx:
        raw_text = extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    # Step 2: Clean text
    raw_text = clean_text(raw_text)

    # Step 3: LLM structured extraction (TODO: M2)
    # For now, return raw_text and a minimal placeholder content_json
    content_json = {
        "schema_version": 1,
        "basic_info": {"name": "", "phone": "", "email": "", "education": "", "school": "", "major": ""},
        "skills": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "self_evaluation": "",
    }

    return raw_text, content_json


def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")


def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


def clean_text(text: str) -> str:
    """Clean extracted text: remove headers/footers/special chars."""
    import re
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove common header/footer patterns
    text = re.sub(r'第\s*\d+\s*页', '', text)
    text = re.sub(r'Page\s*\d+', '', text)
    # Strip
    text = text.strip()
    return text
